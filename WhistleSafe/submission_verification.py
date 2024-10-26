import streamlit as st
import sqlite3
from pathlib import Path
from dfpipeline import analyze_video
from watsonxprocessing import process_question_with_doc, setup_watsonx
import os
from docx import Document

# Add custom CSS for blue theme
st.markdown("""
    <style>
    .stApp {
        background: linear-gradient(to bottom right, #1e3c72, #2a5298);  /* Deep blue gradient background */
    }
    .css-1d391kg, .css-12oz5g7 {
        background-color: rgba(255, 255, 255, 0.9);  /* Semi-transparent white containers */
        padding: 20px;
        border-radius: 10px;
        margin-bottom: 20px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .stButton button {
        background-color: #2e5cb8;  /* Blue buttons */
        color: white;
        border: none;
        border-radius: 5px;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    .stButton button:hover {
        background-color: #3d7be3;  /* Lighter blue on hover */
    }
    h1, h2, h3 {
        color: white !important;
    }
    .st-text {
        color: white;
    }
    </style>
""", unsafe_allow_html=True)

# Add this custom CSS at the beginning
st.markdown("""
    <style>
    .stApp {
        background-color: #f0f8ff;  # Light blue background
    }
    .css-1d391kg, .css-12oz5g7 {
        background-color: white;
        padding: 20px;
        border-radius: 10px;
        margin-bottom: 20px;
    }
    .stButton button {
        background-color: #4CAF50;
        color: white;
    }
    </style>
""", unsafe_allow_html=True)

# Database setup
@st.cache_resource
def get_database_connection():
    conn = sqlite3.connect('user_data.db', check_same_thread=False)
    return conn

conn = get_database_connection()
cursor = conn.cursor()

# Get all reports
def get_all_reports():
    cursor.execute('SELECT id, user_id, video_path, text_report, status FROM uploads')
    return cursor.fetchall()

# Update report status
def update_status(report_id, status):
    cursor.execute('UPDATE uploads SET status = ? WHERE id = ?', (status, report_id))
    conn.commit()

# Verify status update
def verify_status(report_id):
    cursor.execute('SELECT status FROM uploads WHERE id = ?', (report_id,))
    return cursor.fetchone()[0]

# Generate report document
def generate_report(report_id, video_analysis_results, text_analysis_results, status):
    file_path = f'reports/report_{report_id}.docx'
    doc = Document()
    
    # Add report content
    doc.add_heading(f'Report ID: {report_id}', level=1)
    
    if video_analysis_results:
        doc.add_heading('Video Analysis Results:', level=2)
        doc.add_paragraph(str(video_analysis_results))
    
    if text_analysis_results:
        doc.add_heading('Text Analysis Results:', level=2)
        doc.add_paragraph(str(text_analysis_results))
    
    doc.add_heading('Final Status:', level=2)
    doc.add_paragraph(status)
    
    # Save document
    doc.save(file_path)
    return file_path

# Initialize session state for storing analysis results
def init_session_state(report_id):
    if f'video_analysis_{report_id}' not in st.session_state:
        st.session_state[f'video_analysis_{report_id}'] = None
    if f'text_analysis_{report_id}' not in st.session_state:
        st.session_state[f'text_analysis_{report_id}'] = None

# Ensure the reports directory exists
os.makedirs('reports', exist_ok=True)

# Streamlit app layout
st.title("WhistleSafe : Authority Dashboard")

# Initialize WatsonX
watsonx_llm, msg = setup_watsonx()
if not watsonx_llm:
    st.error("Failed to initialize WatsonX. Please check your configuration.")

# Fetch all reports
reports = get_all_reports()

if reports:
    for report in reports:
        report_id, user_id, video_path, text_report, current_status = report
        
        # Initialize session state for this report
        init_session_state(report_id)
        
        # Create a container for each report
        with st.container():
            st.write(f"### Report ID: {report_id}")
            st.write(f"User ID: {user_id}")
            
            # Display video if path exists and is valid
            if os.path.exists(video_path):
                st.video(video_path)
            else:
                st.warning(f"Video file not found: {video_path}")
            
            # Display text report
            st.text_area("Text Report", text_report, height=100, key=f"text_{report_id}")
            st.write(f"Current Status: {current_status}")
            
            # Analysis section
            st.write("### Analysis Tools")
            
            # Video Analysis
            col_video, col_text = st.columns(2)
            with col_video:
                if st.button(f"Analyze Video Report {report_id}"):
                    with st.spinner("Analyzing video..."):
                        try:
                            st.session_state[f'video_analysis_{report_id}'] = analyze_video(video_path)
                            st.success("Video analysis complete!")
                        except Exception as e:
                            st.error(f"Error analyzing video: {str(e)}")
                
                if st.session_state[f'video_analysis_{report_id}']:
                    st.write("Video Analysis Results:")
                    st.write(st.session_state[f'video_analysis_{report_id}'])
            
            # Text Analysis
            with col_text:
                if st.button(f"Analyze Text Report {report_id}"):
                    with st.spinner("Analyzing text..."):
                        if watsonx_llm:
                            try:
                                st.session_state[f'text_analysis_{report_id}'] = process_question_with_doc(text_report, watsonx_llm)
                                st.success("Text analysis complete!")
                            except Exception as e:
                                st.error(f"Error analyzing text: {str(e)}")
                        else:
                            st.error("WatsonX LLM not initialized")
                
                if st.session_state[f'text_analysis_{report_id}']:
                    st.write("Text Analysis Results:")
                    st.write(st.session_state[f'text_analysis_{report_id}'])
            
            # Decision section
            st.write("### Make Decision")
            col1, col2 = st.columns(2)
            
            # Accept button
            with col1:
                if st.button(f"Accept Report {report_id}"):
                    with st.spinner("Processing acceptance..."):
                        update_status(report_id, 'Accepted')
                        updated_status = verify_status(report_id)
                        
                        if updated_status == 'Accepted':
                            file_path = generate_report(
                                report_id,
                                st.session_state[f'video_analysis_{report_id}'],
                                st.session_state[f'text_analysis_{report_id}'],
                                updated_status
                            )
                            st.success(f"Report {report_id} has been accepted. Status confirmed in database.")
                            
                            # Add download button
                            with open(file_path, "rb") as f:
                                st.download_button(
                                    label="Download Accepted Report",
                                    data=f,
                                    file_name=f'accepted_report_{report_id}.docx',
                                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                    key=f"download_accepted_{report_id}"
                                )
                        else:
                            st.error("Status update failed. Please try again.")
            
            # Reject button
            with col2:
                if st.button(f"Reject Report {report_id}"):
                    with st.spinner("Processing rejection..."):
                        update_status(report_id, 'Rejected')
                        updated_status = verify_status(report_id)
                        
                        if updated_status == 'Rejected':
                            file_path = generate_report(
                                report_id,
                                st.session_state[f'video_analysis_{report_id}'],
                                st.session_state[f'text_analysis_{report_id}'],
                                updated_status
                            )
                            st.error(f"Report {report_id} has been rejected. Status confirmed in database.")
                            
                            # Add download button
                            with open(file_path, "rb") as f:
                                st.download_button(
                                    label="Download Rejected Report",
                                    data=f,
                                    file_name=f'rejected_report_{report_id}.docx',
                                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                    key=f"download_rejected_{report_id}"
                                )
                        else:
                            st.error("Status update failed. Please try again.")
            
            st.divider()  # Add visual separation between reports

else:
    st.info("No reports available for review.")

# Cleanup connection when the script ends
def cleanup():
    conn.close()

# Register the cleanup function
import atexit
atexit.register(cleanup)